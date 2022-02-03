import streamlit as st
import streamlit.components as stc
import pandas as pd
import numpy as np
import pathlib
import base64

#import plotly.express as px
#from plotly.subplots import make_subplots
#import plotly.graph_objects as go
#import matplotlib.pyplot as plt

import requests
import urllib
import pandas as pd
from requests_html import HTML
from requests_html import HTMLSession
import trafilatura
import altair as alt
from PIL import Image
from pathlib import Path
import time
#import io
import os
from docx import Document
from fake_useragent import UserAgent
from bs4 import BeautifulSoup
import re

def img_to_bytes(img_path):
    img_bytes = Path(img_path).read_bytes()
    encoded = base64.b64encode(img_bytes).decode()
    return encoded

header_html = "<img src='data:image/png;base64,{}' class='img-fluid'>".format(
    img_to_bytes("DSLogo.png")
)
st.markdown(
    header_html, unsafe_allow_html=True,
)
#image = Image.open('DSLogo.png')

#st.image(image)

#LOGO_IMAGE = "DSLogo.png"
#
#st.markdown(
#    """
#    <style>
#    .container {
#        display: flex;
#    }
#    .logo-text {
#        font-weight:70 !important;
#        font-size:10px !important;
#        color: #f9a01b !important;
#        padding-top: 75px !important;
#    }
#    .logo-img {
#        float:right;
#    }
#    </style>
#    """,
#    unsafe_allow_html=True
#)
#
#st.markdown(
#    f"""
#    <div class="container">
#        <img class="logo-img" src="data:image/png;base64,{base64.b64encode(open(LOGO_IMAGE, "rb").read()).decode()}">
#        <p class="logo-text">Logo Much ?</p>
#    </div>
#    """,
#    unsafe_allow_html=True
#)


st.title("Content Creation for the Given Topic using **_Web Scraping_** and **_NLP_**")
st.sidebar.title("Content Creation for the Given Topic using Web Scraping and NLP")
#st.markdown("This application is to extract URLs and text content related to the given topic:")
st.markdown("***")
st.sidebar.markdown("This application is to extract URLs and text content for the given topic")
#st.sidebar.color_picker("#0068c9")
#st.sidebar.markdown( """ <style> .sidebar .sidebar-content { background-image: linear-gradient(#0068c9,#0068c9); color: blue; } </style> """, unsafe_allow_html=True, )
#bgcolor2 = st.sidebar.color_picker("Pick a Bckground color")
#height = st.sidebar.slider('Height Size',50,200,50)
#width = st.sidebar.slider("Width Size",50,200,50)
# border = st.slider("Border Radius",10,60,10)
#top_left_border = st.sidebar.number_input('Top Left Border',10,50,10)
#top_right_border = st.sidebar.number_input('Top Right Border',10,50,10)
#bottom_left_border = st.sidebar.number_input('Bottom Left Border',10,50,10)
#bottom_right_border = st.sidebar.number_input('Bottom Right Border',10,50,10)

#border_style = st.sidebar.selectbox("Border Style",["dotted","dashed","solid","double","groove","ridge","inset","outset","none","hidden"])
#border_color = st.sidebar.color_picker("Pick a Border Color","#654FEF")

# st.markdown(html_design.format(height,width,bgcolor2,top_left_border,top_right_border,bottom_left_border,bottom_right_border,border_style,border_color),unsafe_allow_html=True)

def get_source(url):
    """Return the source code for the provided URL. 

    Args: 
        url (string): URL of the page to scrape.

    Returns:
        response (object): HTTP response object from requests_html. 
    """

    try:
        session = HTMLSession()
        response = session.get(url)
        return response

    except requests.exceptions.RequestException as e:
        print(e)
        
def scrape_google(query):

    query = urllib.parse.quote_plus(query)
    response = get_source("https://www.google.co.in/search?q=" + query)

    links = list(response.html.absolute_links)
    google_domains = ('https://www.google.', 
                      'https://google.', 
                      'https://webcache.googleusercontent.', 
                      'http://webcache.googleusercontent.', 
                      'https://policies.google.',
                      'https://support.google.',
                      'https://maps.google.',
                      'https://www.youtube.')

    for url in links[:]:
        if url.startswith(google_domains):
            links.remove(url)

    return links

def text_downloader(raw_text):
	b64 = base64.b64encode(raw_text.encode()).decode()
	new_filename = "new_text_file_{}_.docx".format(timestr)
	st.markdown("#### Download File ###")
	href = f'<a href="data:file/docx;base64,{b64}" download="{new_filename}">Click Here!!</a>'
	st.markdown(href,unsafe_allow_html=True)
    
class FileDownloader(object):
	"""docstring for FileDownloader
	>>> download = FileDownloader(data,filename,file_ext).download()
	"""
	def __init__(self, data,filename='myfile',file_ext='docx'):
		super(FileDownloader, self).__init__()
		self.data = data
		self.filename = filename
		self.file_ext = file_ext

	def download(self):
		b64 = base64.b64encode(self.data.encode()).decode()
		new_filename = "{}_{}_.{}".format(self.filename,timestr,self.file_ext)
		st.markdown("#### Download File ###")
		href = f'<a href="data:file/{self.file_ext};base64,{b64}" download="{new_filename}">Click Here!!</a>'
		st.markdown(href,unsafe_allow_html=True)

def download_link(object_to_download, download_filename, download_link_text):
    """
    Generates a link to download the given object_to_download.
    object_to_download (str, pd.DataFrame):  The object to be downloaded.
    download_filename (str): filename and extension of file. e.g. mydata.csv, some_txt_output.txt
    download_link_text (str): Text to display for download link.
    Examples:
    download_link(YOUR_DF, 'YOUR_DF.csv', 'Click here to download data!','required_question_type')
    download_link(YOUR_STRING, 'YOUR_STRING.txt', 'Click here to download your text!','required_question_type')
    """
#     if isinstance(object_to_download,pd.DataFrame):
#         object_to_download = object_to_download.to_csv(index=False)
    doc = Document()
#    with open("C:\\Users\\Darcey\\Downloads\\model_output.txt", 'r', encoding='utf-8') as file:
    doc.add_paragraph(object_to_download)
    doc_to_save = doc.save(str(Topic)+".docx")
#    object_to_download1 = ""
#    if quest_type == "Input Text":
#        object_to_download1 = "="*100+"\r\n"
#    else:
#        object_to_download1 = "-"*100+"\r\n"
#    dt = dtime()
#    object_to_download1+=f"{dt} {quest_type}:\r\n"
#    object_to_download1+="-"*100+"\r\n\r\n"
#    if quest_type == "Match the Following":
#        # object_to_download1+=f"{str(object_to_download)}\r\n"
#        object_to_download1+=tabulate(object_to_download,showindex=False,\
#             headers=object_to_download.columns,tablefmt="grid")
#    elif quest_type == "Input Text":
#            object_to_download1+=f"{object_to_download}"
#    elif quest_type == "Fill in The Blanks":
#        for i,sent in enumerate(object_to_download["sentences"]):
#            object_to_download1+=f"{str(i+1)}. {sent}\r\n"
#        object_to_download1+="\n"+str(object_to_download["keys"])+"\r\n"
#    elif quest_type == "MCQ":
#        count = 1
#        for quest,options in object_to_download.items():
#            asci = 97
#            object_to_download1+=f"{str(count)}. {quest}\n"
#            if options:
#                for opt in options:
#                    object_to_download1+=chr(asci)+")"+" "+opt.title()+"\n"
#                    asci += 1
#                object_to_download1+=f"\r\n"
#            count += 1
#    else:
#        for i,que in enumerate(object_to_download):
#            object_to_download1+=f"{str(i+1)}. {que}\r\n"
#    object_to_download1+=f"\r\n"
#    b64 = base64.b64encode(object_to_download)#.encode()).decode()
    
#    return f'<a href="data:file/docx;base64,{b64}" download="{download_filename}">{download_link_text}</a>'
    return doc_to_save

#Topic = st.text_input('Input the topic here:')

def scrape_google_all(Topic):

    Topic = urllib.parse.quote_plus(Topic) # Format into URL encoding
    number_result = 11

    ua = UserAgent()

    google_url = "https://www.google.com/search?q=" + Topic + "&num=" + str(number_result)
    response = requests.get(google_url, {"User-Agent": ua.random})
    soup = BeautifulSoup(response.text, "html.parser")

    result_div = soup.find_all('div', attrs = {'class': 'ZINbbc'})

    links = []
    titles = []
    descriptions = []
    for r in result_div:
        # Checks if each element is present, else, raise exception
        try:
            link = r.find('a', href = True)
            title = r.find('div', attrs={'class':'vvjwJb'}).get_text()
            description = r.find('div', attrs={'class':'s3v9rd'}).get_text()

            # Check to make sure everything is present before appending
            if link != '' and title != '' and description != '': 
                links.append(link['href'])
                titles.append(title)
                descriptions.append(description)
        # Next loop if one element is not present
        except:
            continue
    return links

def Extract_Ranked_urls(links):
    to_remove = []
    clean_links = []
    for i, l in enumerate(links):
        clean = re.search('\/url\?q\=(.*)\&sa',l)

        # Anything that doesn't fit the above pattern will be removed
        if clean is None:
            to_remove.append(i)
            continue
        clean_links.append(clean.group(1))

    # Remove the corresponding titles & descriptions
#    for x in to_remove:
#        del titles[x]
#        del descriptions[x]
    return clean_links

def Extract_urls(Topic):
    df = pd.DataFrame(scrape_google(Topic), columns = ['link'])
    return df

def Extract_Contents(clean_links):
    list2 = []
    for url in clean_links:
        downloaded = trafilatura.fetch_url(url)
        trafilatura.extract(downloaded)
        # outputs main content and comments as plain text ...
        list1 = trafilatura.extract(downloaded, include_comments=False)
        # outputs main content without comments as XML ...
        list2.append("\n")
        list2.append("---------------------------------------------------------------------------------------------------------------------")
        list2.append("\n")
        list2.append("Below contents are extracted from this url:")
        list2.append("\n")
        list2.append(url)
        list2.append("\n")
        list2.append(list1)
        list3 = ''.join(filter(None, list2))
    return list3

def View_Extracted_Contents(list3):
    Extracted_Contents = list3
#    data = [content.strip() for content in Extracted_Contents.splitlines() if content]
#    data1 = ''.join([str(elem) for elem in data])
    return Extracted_Contents

def para_correct(list3):
    data = [content.strip() for content in list3.splitlines() if content]
    data1 = ''.join([str(elem) for elem in data])
    return data1

def main():
    Topic = st.text_input('Input the topic here and press ENTER:')
    
#if len(Topic)>0:
    if st.sidebar.button("Extract URLs for the given topic"):
        with st.spinner("Extracting..."):
            links = scrape_google_all(Topic)
            clean_links = Extract_Ranked_urls(links)
            st.write("Below are the top URLs to extract content:")
            for x in clean_links:
                st.write(x)
    st.sidebar.markdown("*******************************")
    if st.sidebar.button("Download Contents from URLs"):
#    if text is not None:
        with st.spinner("Downloading..."):
            links = scrape_google_all(Topic)
            clean_links = Extract_Ranked_urls(links)
            list3 = Extract_Contents(clean_links)
#            data1 = para_correct(list3)
            data = [content.strip() for content in list3.splitlines() if content]
            data1 = '\\n\n'.join(f"{row}\n" for row in data)
            doc = Document()
            doc.add_paragraph(data1)
#            docx = Document(io.BytesIO(requests.get(doc).content))
#            b64 = base64.b64encode(docx)  # some strings <-> bytes conversions necessary here
#            href = f'<a href="data:file/docx;base64,{b64}">Download docx file</a>'
#            st.markdown(href, unsafe_allow_html=True)
            doc.paragraph_format.space_after = Inches(1.0)
            doc.save(str(Topic)+".docx")
        st.markdown("Download Complete")
    st.sidebar.markdown("*******************************")
    if st.sidebar.checkbox("View the Extracted Contents"):
        with st.spinner("Downloading the Contents..."):
            links = scrape_google_all(Topic)
            clean_links = Extract_Ranked_urls(links)
            list3 = Extract_Contents(clean_links)
            Extracted_Contents = View_Extracted_Contents(list3)
            data = [content.strip() for content in Extracted_Contents.splitlines() if content]
            for x in data:
                st.write(x)
            list2 = []
            for url in clean_links:
                downloaded = trafilatura.fetch_url(url)
                trafilatura.extract(downloaded)
                # outputs main content and comments as plain text ...
                list1 = trafilatura.extract(downloaded, include_comments=False)
                st.write(url)
                if list1 is None:
                    st.write("Contents not available")
                else:
                    st.write("Contents available")
                ua = UserAgent()
                response = requests.get(url, {"User-Agent": ua.random})

                st.write("Response Code: ", response.status_code)
#        if not st.sidebar.checkbox("View the Extracted Contents"):
#            with st.spinner("Fetching the link to download..."):
#            df = Extract_urls(Topic)
#            list3 = Extract_Contents(df)
#            st.markdown(download_link(list3, 'model_output.txt', 'Click here to download the extracted text'),unsafe_allow_html=True)
#        View_Option = st.sidebar.radio("To view or Download Content: ",
#                     ('View', 'Download'))
#        if (View_Option == 'View'):
#            st.write(list3)
#        elif (View_Option == 'Download'):
#            st.markdown(download_link(list3, 'model_output.txt', 'Click here to download the extracted text'),unsafe_allow_html=True)
main()
