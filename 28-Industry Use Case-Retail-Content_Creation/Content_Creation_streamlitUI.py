#!/usr/bin/env python
# coding: utf-8

# In[56]:


import streamlit as st
import streamlit.components as stc
import pandas as pd
import numpy as np
import pathlib
import base64


import requests
import urllib
import pandas as pd
from requests_html import HTML
from requests_html import HTMLSession
import trafilatura
import altair as alt
from PIL import Image
from pathlib import Path
import time
import io
import os
import docx
from docx import Document
from fake_useragent import UserAgent
from bs4 import BeautifulSoup
import re

from docx.shared import Inches, Cm
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_UNDERLINE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, ns
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
import nltk
nltk.download('punkt')
from nltk.tokenize import sent_tokenize

from urllib.request import Request, urlopen
from bs4 import BeautifulSoup as soup
import mock

import docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import unicodedata


def img_to_bytes(img_path):
    img_bytes = Path(img_path).read_bytes()
    encoded = base64.b64encode(img_bytes).decode()
    return encoded

header_html = "<img src='data:image/png;base64,{}' class='img-fluid'>".format(
    img_to_bytes("C:\\Users\\Darcey\\Downloads\\DeepSphere Logo.jpg")
)
st.markdown(
    header_html, unsafe_allow_html=True,
)


st.title("Content Creation for the Given Topic using **_Web Scraping_** and **_NLP_**")
st.sidebar.title("Content Creation for the Given Topic using Web Scraping and NLP")
#st.markdown("This application is to extract URLs and text content related to the given topic:")
st.markdown("***")
st.sidebar.markdown("This application is to extract URLs and text content for the given topic")

def get_source(url):
    """Return the source code for the provided URL. 

    Args: 
        url (string): URL of the page to scrape.

    Returns:
        response (object): HTTP response object from requests_html. 
    """

    try:
        session = HTMLSession()
        response = session.get(url)
        return response

    except requests.exceptions.RequestException as e:
        print(e)
        
def scrape_google(query):

    query = urllib.parse.quote_plus(query)
    response = get_source("https://www.google.co.in/search?q=" + query)

    links = list(response.html.absolute_links)
    google_domains = ('https://www.google.', 
                      'https://google.', 
                      'https://webcache.googleusercontent.', 
                      'http://webcache.googleusercontent.', 
                      'https://policies.google.',
                      'https://support.google.',
                      'https://maps.google.',
                      'https://www.youtube.')

    for url in links[:]:
        if url.startswith(google_domains):
            links.remove(url)

    return links

def text_downloader(raw_text):
	b64 = base64.b64encode(raw_text.encode()).decode()
	new_filename = "new_text_file_{}_.docx".format(timestr)
	st.markdown("#### Download File ###")
	href = f'<a href="data:file/docx;base64,{b64}" download="{new_filename}">Click Here!!</a>'
	st.markdown(href,unsafe_allow_html=True)
    
class FileDownloader(object):
	"""docstring for FileDownloader
	>>> download = FileDownloader(data,filename,file_ext).download()
	"""
	def __init__(self, data,filename='myfile',file_ext='docx'):
		super(FileDownloader, self).__init__()
		self.data = data
		self.filename = filename
		self.file_ext = file_ext

	def download(self):
		b64 = base64.b64encode(self.data.encode()).decode()
		new_filename = "{}_{}_.{}".format(self.filename,timestr,self.file_ext)
		st.markdown("#### Download File ###")
		href = f'<a href="data:file/{self.file_ext};base64,{b64}" download="{new_filename}">Click Here!!</a>'
		st.markdown(href,unsafe_allow_html=True)

def download_link(object_to_download, download_filename, download_link_text):
    """
    Generates a link to download the given object_to_download.
    object_to_download (str, pd.DataFrame):  The object to be downloaded.
    download_filename (str): filename and extension of file. e.g. mydata.csv, some_txt_output.txt
    download_link_text (str): Text to display for download link.
    Examples:
    download_link(YOUR_DF, 'YOUR_DF.csv', 'Click here to download data!','required_question_type')
    download_link(YOUR_STRING, 'YOUR_STRING.txt', 'Click here to download your text!','required_question_type')
    """
#     if isinstance(object_to_download,pd.DataFrame):
#         object_to_download = object_to_download.to_csv(index=False)
    doc = Document()
#    with open("C:\\Users\\Darcey\\Downloads\\model_output.txt", 'r', encoding='utf-8') as file:
    doc.add_paragraph(object_to_download)
    doc_to_save = doc.save(str(Topic)+".docx")
#    object_to_download1 = ""

    return doc_to_save

#Topic = st.text_input('Input the topic here:')

def scrape_google_all(Topic):

    Topic = urllib.parse.quote_plus(Topic) # Format into URL encoding
    number_result = 10

    ua = UserAgent()

    google_url = "https://www.google.com/search?q=" + Topic + "&num=" + str(number_result)
    response = requests.get(google_url, {"User-Agent": ua.random})
    soup = BeautifulSoup(response.text, "html.parser")

    result_div = soup.find_all('div', attrs = {'class': 'ZINbbc'})

    links = []
    titles = []
    descriptions = []
    for r in result_div:
        # Checks if each element is present, else, raise exception
        try:
            link = r.find('a', href = True)
            title = r.find('div', attrs={'class':'vvjwJb'}).get_text()
            description = r.find('div', attrs={'class':'s3v9rd'}).get_text()

            # Check to make sure everything is present before appending
            if link != '' and title != '' and description != '': 
                links.append(link['href'])
                titles.append(title)
                descriptions.append(description)
        # Next loop if one element is not present
        except:
            continue
    google_domains = ('https://www.google.', 
                      'https://google.', 
                      'https://webcache.googleusercontent.', 
                      'http://webcache.googleusercontent.', 
                      'https://policies.google.',
                      'https://support.google.',
                      'https://maps.google.',
                      'https://www.youtube.')

    for url in links[:]:
        if url.startswith(google_domains):
            links.remove(url)
    return links

def Extract_Ranked_urls(links):
    to_remove = []
    clean_links = []
    for i, l in enumerate(links):
        clean = re.search('\/url\?q\=(.*)\&sa',l)

        # Anything that doesn't fit the above pattern will be removed
        if clean is None:
            to_remove.append(i)
            continue
        clean_links.append(clean.group(1))

    # Remove the corresponding titles & descriptions
#    for x in to_remove:
#        del titles[x]
#        del descriptions[x]
    return clean_links

def Extract_urls(Topic):
    df = pd.DataFrame(scrape_google(Topic), columns = ['link'])
    return df



def View_Extracted_Contents(list3):
    Extracted_Contents = list3
#    data = [content.strip() for content in Extracted_Contents.splitlines() if content]
#    data1 = ''.join([str(elem) for elem in data])
    return Extracted_Contents

def para_correct(list3):
    data = [content.strip() for content in list3.splitlines() if content]
    data1 = ''.join([str(elem) for elem in data])
    return data1

##################### To add page number in the footer ###############################################

def create_element(name):
    return OxmlElement(name)


def create_attribute(element, name, value):
    element.set(ns.qn(name), value)


def add_page_number(paragraph):
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    page_run = paragraph.add_run()
    t1 = create_element('w:t')
    create_attribute(t1, 'xml:space', 'preserve')
    t1.text = 'Page '
    page_run._r.append(t1)

    page_num_run = paragraph.add_run()

    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    page_num_run._r.append(fldChar1)
    page_num_run._r.append(instrText)
    page_num_run._r.append(fldChar2)

    of_run = paragraph.add_run()
    t2 = create_element('w:t')
    create_attribute(t2, 'xml:space', 'preserve')
    t2.text = ' of '
    of_run._r.append(t2)

    fldChar3 = create_element('w:fldChar')
    create_attribute(fldChar3, 'w:fldCharType', 'begin')

    instrText2 = create_element('w:instrText')
    create_attribute(instrText2, 'xml:space', 'preserve')
    instrText2.text = "NUMPAGES"

    fldChar4 = create_element('w:fldChar')
    create_attribute(fldChar4, 'w:fldCharType', 'end')

    num_pages_run = paragraph.add_run()
    num_pages_run._r.append(fldChar3)
    num_pages_run._r.append(instrText2)
    num_pages_run._r.append(fldChar4)
    
##################### To add page number in the footer ###############################################

def remove_control_characters(s):
    return "".join(ch for ch in s if unicodedata.category(ch)[0]!="C")

def main():
    Topic = st.text_input('Input the topic here and press ENTER:')
    i=1
    #if len(Topic)>0:
    if st.sidebar.button("Extract URLs for the given topic"):
        with st.spinner("Extracting..."):
            links = scrape_google_all(Topic)
            clean_links = Extract_Ranked_urls(links)
            st.write("Below are the top URLs to extract content:")
            for x in clean_links:
                st.write(x)
    st.sidebar.markdown("*******************************")
    if st.sidebar.button("Download Contents from URLs"):
        with st.spinner("Downloading..."):
            links = scrape_google_all(Topic)
            clean_links = Extract_Ranked_urls(links)

############ Converting listof urls to dataframe and then to tuple to create the table in word document##########

            df = pd.DataFrame(clean_links, columns = ['urls'])
            df['url_rank'] = np.arange(len(df)) + 1
            df1 = df[['url_rank', 'urls']]
            Search_for_These_values = ['youtube','.pdf','.pptx'] 
            pattern = '|'.join(Search_for_These_values)
            df2 = df1.loc[~df1['urls'].str.contains(pattern, case=False)]
            #            print (df1)
            datat = tuple(df1.itertuples(index=False, name=None))


            #            print(datat)
            doc = docx.Document()

##################### To add table of contents ###############################################


            paragraph = doc.add_paragraph()
            toc = paragraph.add_run("\t Table of Contents - " + str(Topic).upper())
            toc.bold = True
            toc.font.size = Pt(14)
            toc.font.color.rgb = RGBColor(0, 0, 0)
            run = paragraph.add_run()
            fldChar = OxmlElement('w:fldChar')  # creates a new element
            fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
            instrText = OxmlElement('w:instrText')
            instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
            instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need

            fldChar2 = OxmlElement('w:fldChar')
            fldChar2.set(qn('w:fldCharType'), 'separate')
            fldChar3 = OxmlElement('w:updateFields')
            fldChar3.set(qn('w:val'), 'true')
            #fldChar3.text = "Right-click to update field."
            fldChar2.append(fldChar3)

            fldChar4 = OxmlElement('w:fldChar')
            fldChar4.set(qn('w:fldCharType'), 'end')

            r_element = run._r
            r_element.append(fldChar)
            r_element.append(instrText)
            r_element.append(fldChar2)
            r_element.append(fldChar4)
            p_element = paragraph._p

            doc.add_page_break()

##################### To add logo ###############################################

        ##### add logo in Zoned header

            logo_path = 'C:\\Users\\Darcey\\Downloads\\DeepSphere Logo.jpg'    # Path of the image file
            section = doc.sections[0]   # Create a section
            sec_header = section.header   # Create header 
            header_tp = sec_header.add_paragraph()  # Add a paragraph in the header, you can add any anything in the paragraph
            header_run = header_tp.add_run()   # Add a run in the paragraph. In the run you can set the values 
            header_run.add_picture(logo_path, width=Inches(1.3))  # Add a picture and set width.
            #rml_header = "\t Applied Artificial Intelligence for Schools Content \t Generation by Topic \t"
            header_run.add_text("\n                                                                                                        ")
            header_run.add_text("Applied Artificial Intelligence for Schools Content Generation by Topic")
            header_run.add_text("\n_______________________________________________________________________________________")
            header_run.font.size =  Pt(13)
            header_run.font.color.rgb = RGBColor(0, 0, 0)
            header_run.font.bold = True


            doc.add_paragraph('')
#            doc.add_paragraph('')

##################### To add footer with page number ###############################################

            section = doc.sections[0]
            footer = section.footer
            footer_para = footer.paragraphs[0]
            footer_para.text = "_________________________________________________________________________________ \t \n\n © DeepSphere.AI | Confidential and Proprietary |Not for Distribution \t"
            add_page_number(doc.sections[0].footer.paragraphs[0])

##################### To add table with url and its ranking ###############################################

#Add heading for the table

            table_heading = doc.add_heading("Extracted URls and its Ranking for the Given Topic:", 1)
            doc.add_paragraph('')
            table_heading.style.font.color.rgb = RGBColor(0, 0, 153)
            table_heading.style.font.size = Pt(16)
            table_heading.style.font.bold = True
            table_heading.style.font.all_caps = True

            table = doc.add_table(rows=1, cols=2)
            row = table.rows[0].cells
            row[0].text = 'URL Rank'
            row[1].text = 'URLs'
            for url_rank, urls in datat:
                row = table.add_row().cells
                row[0].text = str(url_rank) + "                               "
                row[1].text = urls

            table.style = 'Colorful List'
            table.autofit = False
            table.allow_autofit = False
            cell = table.rows[0].cells[0]
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell1 = table.rows[0].cells[1]
            cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

##################### To export/save the document ###############################################


            #            for x in list2:

            #    doc = docx.Document()
            doc.add_page_break()
            #            text_split = []
            links = scrape_google_all(Topic)
            clean_links = Extract_Ranked_urls(links)
            ua = UserAgent()
            i=1
            for url in clean_links:
                para1 = doc.add_paragraph().add_run("-----------------------------------------------------------------------------------------")
                para1.bold = True
                para1.font.size = Pt(14)
                para1.font.color.rgb = RGBColor(0, 0, 153)
            #    docu.add_paragraph("")
                txt1 = "**Content Set #" + str(i) + "**"
                para2 = doc.add_paragraph().add_run(txt1)
                para2.bold = True
                para2.font.size = Pt(14)
                para2.font.color.rgb = RGBColor(0, 0, 153)
            #    docu.add_paragraph("")
            #    docu.add_paragraph("")
                txt2 = "URL #" + str(i) + ":    " + url
                para3 = doc.add_paragraph().add_run(txt2)
                para3.bold = True
                para3.font.size = Pt(14)
                para3.font.color.rgb = RGBColor(0, 0, 153)
            #    docu.add_paragraph("")
                para4 = doc.add_paragraph().add_run("-----------------------------------------------------------------------------------------")
                para4.bold=True
                para4.font.size = Pt(14)
                para4.font.color.rgb = RGBColor(0, 0, 153)

                with mock.patch.object(requests.cookies.RequestsCookieJar, 'update', lambda *args, **kwargs: 0):
                    req = Request(url , headers={'User-Agent': ua.random})
                    req.add_header('User-Agent', 'Mozilla/4.0 (compatible; MSIE 7.0; Windows NT 6.0)')
                    try:

                        webpage = urlopen(req).read()

                    except Exception as e:
                        err = "Exception occured while extracting content - " + str(e)
                        doc.add_paragraph(err)
                    else:
                        page_soup = soup(webpage, "html.parser")
                        heading_tags = ["h1", "h2", "h3", "h4"]
                        list2 = []
                        for tags in page_soup.find_all([heading_tags, 'p']):
                            if not tags.find([heading_tags, 'p']):
                                list1 = tags.name + ' -> ' + tags.text.strip()
                                if list1.startswith("h1"):
                                    h1 = doc.add_heading(list1.replace("h1 -> ", ""), 1)
                                    h1.style.font.color.rgb = RGBColor(0, 0, 139)
                                    h1.style.font.size = Pt(14)
                                    h1.style.font.bold = True
                                    h1.style.font.all_caps = True
                                elif list1.startswith("h2"):
                                    h2 = doc.add_heading(list1.replace("h2 -> ", ""), 2)
                                    h2.style.font.color.rgb = RGBColor(0, 0, 205)
                                    h2.style.font.size = Pt(12)
                                    h2.style.font.bold = True
                                    h2.style.font.all_caps = True
                                elif list1.startswith("h3"):
                                    h3 = doc.add_heading(list1.replace("h3 -> ", ""), 3)
                                    h3.style.font.color.rgb = RGBColor(0, 0, 255)
                                    h3.style.font.size = Pt(10)
                                    h3.style.font.bold = True
                                    h3.style.font.all_caps = True
                                elif list1.startswith("h4"):
                                    h4 = doc.add_heading(list1.replace("h4 -> ", ""), 4)
                                    h4.style.font.color.rgb = RGBColor(65, 105, 225)
                                    h4.style.font.size = Pt(10)
                                    h4.style.font.bold = True
                                    h4.style.font.all_caps = True
                                else:
                                    try:
                                        print(list1.replace("p -> ", ""))
                                    except Exception as ee:
                                        doc.add_paragraph(str(ee))
                                    else:
                                        doc.add_paragraph(remove_control_characters(list1.replace('\x00','')).replace("p -> ", ""))
                                list2.append("\n\n")
                                list2.append(list1)
                                list3 = ''.join(filter(None, list2))

            #    doc.add_paragraph(list1)
                doc.add_paragraph("")
                para5 = doc.add_paragraph().add_run("------------------------------------------------------------------------------------------")
                para5.bold=True
                para5.font.size = Pt(14)
                para5.font.color.rgb = RGBColor(0, 0, 153)
                #    docu.add_paragraph("")
                para6 = doc.add_paragraph().add_run("------------------------------------------------------------------------------------------")
                para6.bold=True
                para6.font.size = Pt(14)
                para6.font.color.rgb = RGBColor(0, 0, 153)

                doc.add_page_break()
                i+=1

            doc.save("Model Output - " + str(Topic)+".docx")

main()

